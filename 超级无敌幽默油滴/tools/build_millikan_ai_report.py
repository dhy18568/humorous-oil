from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "artifacts" / "documents"
OUT_PATH = OUT_DIR / "密立根油滴AI实验平台分析报告.docx"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
BORDER = "B8C2CC"


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: str = INK):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Calibri"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[idx]))

            tc_mar = tc_pr.find(qn("w:tcMar"))
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side, value in {
                "top": "80",
                "bottom": "80",
                "start": "120",
                "end": "120",
            }.items():
                node = tc_mar.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    tc_mar.append(node)
                node.set(qn("w:w"), value)
                node.set(qn("w:type"), "dxa")


def set_style_font(style, size_pt, color=INK, bold=False, space_after=6, space_before=0, line_spacing=1.10):
    font = style.font
    font.name = "Calibri"
    font.size = Pt(size_pt)
    font.bold = bold
    font.color.rgb = RGBColor.from_string(color)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    pf = style.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing


def add_para(doc, text: str, style=None, bold_prefix: str | None = None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r2 = p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    for run in p.runs:
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.color.rgb = RGBColor.from_string(INK)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        r.font.name = "Calibri"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor.from_string(INK)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        r.font.name = "Calibri"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor.from_string(INK)


def add_callout(doc, title: str, body: str):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Calibri"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    r2.font.name = "Calibri"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor.from_string(INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_matrix(doc, headers, rows, widths, header_fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for idx, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[idx], header_fill)
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color=DARK_BLUE)
    for row_data in rows:
        row = table.add_row()
        for idx, text in enumerate(row_data):
            set_cell_text(row.cells[idx], text)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_code_line(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor.from_string("374151")


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    set_style_font(styles["Normal"], 11, INK, False, 6, 0, 1.10)
    set_style_font(styles["Title"], 22, BLUE, True, 3, 0, 1.0)
    set_style_font(styles["Heading 1"], 16, BLUE, True, 16, 8, 1.10)
    set_style_font(styles["Heading 2"], 13, BLUE, True, 12, 6, 1.10)
    set_style_font(styles["Heading 3"], 12, DARK_BLUE, True, 8, 4, 1.10)
    set_style_font(styles["List Bullet"], 10.5, INK, False, 4, 0, 1.167)
    set_style_font(styles["List Number"], 10.5, INK, False, 4, 0, 1.167)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("HumOil 密立根油滴 AI 实验平台分析报告")
    fr.font.name = "Calibri"
    fr._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor.from_string("6B7280")

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("密立根油滴 AI 实验平台分析报告")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("聚类发现、深度学习辅助拟合与物理实验教学价值")
    sr.font.name = "Calibri"
    sr._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    sr.font.size = Pt(13)
    sr.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run("基于 D:\\GitHub\\humoil 当前实现整理 | 2026 年 6 月")
    mr.font.name = "Calibri"
    mr._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    mr.font.size = Pt(10)
    mr.font.color.rgb = RGBColor.from_string("6B7280")

    add_callout(
        doc,
        "报告结论",
        "这套平台的核心价值不是让 AI 替学生给出一个不可解释答案，而是把密立根油滴实验拆成可观察、可调参、可验证的数据处理链路：先用物理公式把 U、t 转换成连续电荷估计，再用无监督机器学习发现电荷分布中的自然簇，随后用半峰宽筛选高置信点；在拟合阶段，系统用 MLP 集成作为深度学习 teacher 对多条 U-t 曲线进行去噪学习，再把 teacher 学到的平滑关系蒸馏为可解释的符号表达式，并保留传统物理约束拟合作为对照。"
    )

    doc.add_heading("1. 平台定位与教学目标", level=1)
    add_para(doc, "密立根油滴实验的经典目标是从单个油滴的受力平衡和运动过程推断电荷量，并进一步观察电荷量是否呈现离散倍数关系。传统教学中，学生通常依靠表格计算、人工剔除异常值、手工分组和线性或非线性拟合来得到结果。HumOil 平台将这一过程变为一个可交互的数据科学工作流，让学生在同一个实验场景中同时学习物理建模、数据清洗、机器学习聚类、深度学习拟合和可解释模型评估。")
    add_para(doc, "平台当前的核心路径位于 `src/humoil/core/oil/tabs/regression.py`、`tab_classify.py` 和 `tab_regress.py`。它保留传统物理拟合路径，但主线已经转向“AI 发现式聚类与符号回归”：不预先把数据硬分配到某个整数电荷倍数，而是先从测量数据中发现 Q 分布结构，再用拟合和对照检验解释该结构。")

    add_matrix(
        doc,
        ["教学维度", "平台实现", "学生应形成的能力"],
        [
            ["物理建模", "由下落时间 t、平衡电压 U、油滴密度、空气密度、黏滞系数、极板间距等参数计算电荷估计 Q。", "理解实验读数不是结论；必须经过物理公式和量纲转换才能成为可分析变量。"],
            ["AI 聚类", "支持 K-Means、Gaussian Mixture、DBSCAN 和 KDE 峰发现，从 Q 分布中识别自然簇。", "理解无监督学习如何发现数据结构，以及聚类参数会如何影响实验结论。"],
            ["深度学习", "使用多个 MLPRegressor 组成 teacher ensemble，学习平滑的 U=f(t,Q_c) 曲面。", "理解深度学习可以作为去噪和函数逼近工具，但需要被验证和解释。"],
            ["可解释拟合", "将神经网络 teacher 蒸馏为幂律族候选公式，并用 RMSE、MAE、R2、BIC 选择模型。", "学习如何把 AI 输出转化为可讨论的物理表达式，而不是只看黑箱预测。"],
            ["科学验证", "保留传统物理约束拟合作为对照，显示残差、参与拟合点和簇质量。", "形成模型比较、异常点处理、证据链追踪的科学习惯。"],
        ],
        [1700, 3900, 3760],
        header_fill=LIGHT_BLUE,
    )

    doc.add_heading("2. 总体数据流", level=1)
    add_para(doc, "平台的数据处理流程可以概括为七个阶段。每个阶段既对应一个算法步骤，也对应一个学生需要理解的实验方法问题。")
    add_numbers(doc, [
        "录入或加载油滴测量数据：核心输入是 FallingTime(t/s) 和 BalanceVoltage(U/V)。",
        "用物理公式计算连续电荷估计 Q：得到 ChargeEstimate(C) 和 ChargeEstimate(1e-19C)。",
        "在 Q 空间做无监督聚类或密度峰发现：识别候选电荷簇中心 Q_c。",
        "用半峰宽容差标记高置信点：生成 UseForFit、ChargeCluster、ChargeDistance 和 ClusterQuality。",
        "对高置信点执行共享符号回归：搜索 U(t,Q_c) 的候选表达式。",
        "在数据量足够时训练 MLP teacher：用深度学习先拟合平滑曲面，再进行符号蒸馏。",
        "输出图表、候选式、残差和报告：学生可以比较 AI 聚类、符号式和传统物理拟合的一致性。"
    ])

    add_matrix(
        doc,
        ["阶段", "输入", "关键输出", "可解释含义"],
        [
            ["物理特征计算", "t, U, 实验参数", "Velocity, Radius, ChargeEstimate", "把宏观读数转为油滴半径、电荷估计等物理量。"],
            ["AI 聚类", "ChargeEstimate(1e-19C)", "ChargeCluster, Q_center", "让数据自己显示可能的电荷层级，而非先验指定整数倍。"],
            ["半峰宽筛选", "簇中心与容差", "UseForFit, ClusterQuality", "区分高置信点、半宽外点和被禁用簇。"],
            ["深度学习 teacher", "log(t), log(Q_c), U", "TeacherVoltage, uncertainty", "学习平滑函数关系，降低噪声点对符号式的影响。"],
            ["符号蒸馏", "teacher 曲线或高置信点", "U(t,Q_c) 公式、RMSE、R2、BIC", "把 AI 学到的关系压缩成学生可读的数学表达式。"],
        ],
        [1550, 2100, 2500, 3210],
        header_fill=LIGHT_BLUE,
    )

    doc.add_heading("3. 从物理测量到机器学习特征", level=1)
    add_para(doc, "平台没有直接把原始电压和时间丢给神经网络，而是先做物理特征工程。`add_charge_estimates` 函数根据下落距离、油滴与空气密度差、黏滞系数、重力加速度和极板间距，计算油滴速度、半径、有效重力以及电荷估计。")
    add_code_line(doc, "v = fall_distance / t")
    add_code_line(doc, "r = sqrt(9 * viscosity * v / (2 * density_delta * g))")
    add_code_line(doc, "q = effective_weight * plate_distance / U")
    add_para(doc, "这一步非常重要：AI 的输入不只是数据表，而是经过物理约束转换后的变量。学生由此可以看到，AI 在物理实验中常常不是替代物理公式，而是接在物理建模之后，用来处理噪声、分组和复杂关系发现。")

    doc.add_heading("4. AI 聚类逻辑：从 Q 分布中发现自然簇", level=1)
    add_para(doc, "当前平台将“数据分类”重新定位为机器学习聚类分析。它不是用人工标签训练一个分类器去判断某个点属于哪一类，而是在连续电荷估计 Q 的分布上寻找自然簇。学生可以选择 K-Means、Gaussian Mixture、DBSCAN 或 KDE 峰发现。")
    add_matrix(
        doc,
        ["方法", "系统用途", "教学含义"],
        [
            ["K-Means", "按预期簇数把 Q 值分组，适合作为默认快速聚类方法。", "帮助学生理解簇数假设会影响结果。"],
            ["Gaussian Mixture", "用概率混合模型描述可能重叠的 Q 分布。", "展示“属于某簇”可以是概率性的，而不是硬边界。"],
            ["DBSCAN", "根据局部密度识别簇，并允许噪声点存在。", "适合说明异常点和低密度点不应被强行归类。"],
            ["KDE 峰发现", "先估计 Q 分布密度，再找显著峰和半峰宽。", "更贴近“从电荷分布中找峰”的实验直觉。"],
        ],
        [1800, 3600, 3960],
        header_fill=LIGHT_BLUE,
    )
    add_para(doc, "聚类后，系统并不把所有点都纳入拟合。每个 Q 峰会有中心 `ChargeCenter(1e-19C)` 和半峰宽 `ChargeHalfWidth(1e-19C)`。只有距离峰中心不超过半峰宽的点才被标记为 `UseForFit=True`；半宽之外的点保留在明细表里，但标记为 `half_width_outlier`，不参与主公式拟合。")
    add_callout(
        doc,
        "刚刚修复的边界情况",
        "未进入半峰宽的点没有正式的 ChargeCluster，因此该列会出现 NaN。最新修复保证系统只对非空簇编号做整数转换，避免在选择参与拟合簇时因 NaN 转 int 而失败。这体现了实验数据处理中对缺失值和离群点的严谨处理。"
    )

    doc.add_heading("5. 深度学习辅助拟合：MLP teacher 不是黑箱结论，而是去噪中间层", level=1)
    add_para(doc, "在可用高置信点足够多时，平台会训练一个神经网络 teacher。该 teacher 使用 `MLPRegressor`，网络结构为两层隐藏层 `(48, 24)`，激活函数为 `tanh`，并通过多个随机种子训练出集成模型。输入特征不是原始 t 和 Q，而是 `log(t)` 与 `log(Q_c)`，这与幂律关系的物理直觉相匹配。")
    add_bullets(doc, [
        "训练目标：学习 U 与 t、Q_c 之间的平滑函数关系。",
        "集成策略：默认最多训练 5 个 MLP teacher，用平均预测降低单个模型偶然性。",
        "验证方式：使用训练/测试拆分，输出 test RMSE、MAE、R2 和不确定度。",
        "教学边界：神经网络不直接替代物理结论；它先作为去噪函数逼近器，再接受符号回归蒸馏。"
    ])
    add_para(doc, "这种设计很适合教学：学生能够看到深度学习擅长从噪声数据中学习平滑关系，但最终仍需要被压缩成公式、被误差指标检验、被传统物理模型对照。")

    doc.add_heading("6. 符号回归与蒸馏：把 AI 曲面变成可解释公式", level=1)
    add_para(doc, "平台的符号回归不是任意生成公式，而是在有限的候选族中搜索。候选族包括 `single_power`、`additive_power`、`time_correction` 和 `charge_correction`。系统在电荷幂指数和时间幂指数网格上枚举候选，并用最小二乘求系数。")
    add_matrix(
        doc,
        ["候选式族", "直观解释", "适合回答的问题"],
        [
            ["single_power", "一个主要幂律项加偏置。", "是否存在简洁的主规律？"],
            ["additive_power", "时间项与电荷项相加。", "两个因素是否分别贡献电压变化？"],
            ["time_correction", "主项外加时间修正。", "时间测量误差或阻力效应是否显著？"],
            ["charge_correction", "主项外加电荷修正。", "不同电荷簇之间是否存在系统偏差？"],
        ],
        [1900, 3650, 3810],
        header_fill=LIGHT_BLUE,
    )
    add_para(doc, "当 MLP teacher 可用时，平台优先对 teacher 产生的平滑曲线做符号蒸馏。也就是说，神经网络先学习复杂但平滑的函数曲面，符号回归再用少量幂律候选式去逼近它。最终选择不仅看误差，还看公式复杂度：系统用 RMSE、MAE、R2 和 BIC 综合评估，并在误差接近时偏好更简单、物理上更容易解释的表达式。")

    doc.add_heading("7. 两阶段发现与全局候选搜索", level=1)
    add_para(doc, "为了避免单一搜索路径误导学生，平台还保留了两阶段符号发现和全局候选搜索。两阶段路径先让每个 Q 峰拥有独立系数，只共同搜索时间幂指数；随后再把每个峰的系数作为样本，搜索其与 Q_c 的幂关系。")
    add_numbers(doc, [
        "阶段一：在每个电荷簇内部拟合 U 与 t 的关系，比较不同时间幂指数的 RMSE 和 BIC。",
        "阶段二：把各簇的时间系数与 Q_c 建立关系，搜索电荷幂指数。",
        "模型选择：在误差接近的候选中，优先选择幂指数更简单、BIC 更低的公式。",
        "对照作用：如果神经网络 teacher 不可用，系统可回退到两阶段或全局候选搜索。"
    ])
    add_para(doc, "这一设计帮助学生理解 AI 建模不是一次按钮点击，而是多条证据链之间的比较：神经网络能拟合、符号式能解释、两阶段路径能展示结构、传统物理模型能校验方向。")

    doc.add_heading("8. 传统物理拟合作为对照", level=1)
    add_para(doc, "平台保留了 `physics_guided_regression` 路径。它以传统整数 n 的思路工作：清洗数据、估计全局参数 A 和 b、计算浮点 `PhysicsNFloat`、四舍五入为 `PhysicsN`，再按整数峰宽过滤点，并迭代修正参数、剔除残差异常点。")
    add_matrix(
        doc,
        ["对照路径", "AI 发现式路径", "教学价值"],
        [
            ["先验假设电荷整数倍结构。", "先从 Q 分布中发现自然簇，再讨论是否接近等间距。", "区分物理先验和数据发现。"],
            ["重点是拟合 A、b 和整数 n。", "重点是聚类中心 Q_c、半峰宽、teacher 曲线和符号表达式。", "把实验误差、聚类和模型选择显性化。"],
            ["异常点通过整数峰宽和残差剔除。", "异常点通过半峰宽、UseForFit 和 ClusterQuality 标记。", "学习不同异常点定义的影响。"],
        ],
        [2800, 3300, 3260],
        header_fill=LIGHT_BLUE,
    )

    doc.add_heading("9. 学生能够学到的 AI + 物理实验知识技能", level=1)
    add_bullets(doc, [
        "物理量构造能力：理解如何从 t、U 等原始读数计算速度、半径和电荷估计，而不是直接依赖模型。",
        "无监督学习能力：理解 K-Means、GMM、DBSCAN、KDE 峰发现的适用场景和参数含义。",
        "数据质量意识：理解半峰宽筛选、离群点、缺失簇编号和 UseForFit 的意义。",
        "深度学习边界意识：知道 MLP teacher 是去噪和函数逼近工具，不是实验结论本身。",
        "可解释 AI 能力：学习如何把神经网络曲面蒸馏为符号公式，并用误差和复杂度指标选择模型。",
        "科学对照能力：比较 AI 发现式结果与传统物理约束拟合结果，判断二者是否相互支持。",
        "实验报告能力：把图表、参数、拟合结果、残差和异常点处理过程组织成可复查证据链。"
    ])

    doc.add_heading("10. 教学实施建议", level=1)
    add_para(doc, "建议将平台用于三段式教学：第一段让学生只录入数据并观察原始 Q 分布；第二段让学生调节聚类方法、簇数、半峰宽和 KDE 参数，观察参与拟合点如何变化；第三段再执行符号回归，比较 neural teacher、两阶段搜索和传统物理拟合的结果。")
    add_matrix(
        doc,
        ["课堂任务", "操作建议", "讨论问题"],
        [
            ["数据录入与特征观察", "先不显示参考背景，只看学生数据的 Q-U 散点。", "为什么原始 U-t 读数不能直接说明电荷量？"],
            ["聚类参数实验", "比较 K-Means、GMM、DBSCAN、KDE 的簇中心和半峰宽。", "哪些点应被视为异常？参数改变是否改变物理解释？"],
            ["深度学习 teacher", "观察 teacher 曲线、测试 RMSE、R2 和不确定度。", "神经网络拟合得好是否等于物理规律正确？"],
            ["符号公式评估", "查看候选公式、BIC、RMSE、R2、残差图。", "为什么误差最低的公式不一定是最适合教学解释的公式？"],
            ["传统对照", "执行传统物理拟合对照检验。", "AI 发现式簇与整数 n 模型是否一致？差异来自数据、参数还是模型假设？"],
        ],
        [1800, 3600, 3960],
        header_fill=LIGHT_BLUE,
    )

    doc.add_heading("11. 局限性与后续改进方向", level=1)
    add_bullets(doc, [
        "当前 teacher 是表格数据上的 MLP 集成，不是图像端到端识别；视觉测量仍应作为独立模块逐步完善。",
        "聚类结果对参数敏感，教学时应要求学生记录参数，而不是只提交最终图。",
        "符号回归候选族是人为设计的有限搜索空间，适合教学解释，但不代表穷尽所有物理模型。",
        "参考数据背景和大规模图表需要按需显示；当前已通过默认关闭参考背景和抽样绘图改善响应速度。",
        "应继续增加自动化测试，覆盖 NaN 簇编号、空数据、少量数据、单簇数据和高噪声数据等边界情况。"
    ])

    doc.add_heading("附录 A：关键参数与代码位置", level=1)
    add_matrix(
        doc,
        ["项目", "当前实现"],
        [
            ["入口", "streamlit_app.py 调用 src/humoil/core/oil/app.py。"],
            ["AI 聚类页", "src/humoil/core/oil/tabs/tab_classify.py。"],
            ["符号回归页", "src/humoil/core/oil/tabs/tab_regress.py。"],
            ["核心算法", "src/humoil/core/oil/tabs/regression.py。"],
            ["默认聚类", "KMeans，requested_clusters=5，half_width_1e19c=0.25。"],
            ["KDE 参数", "kde_bandwidth=0.08，peak_prominence=0.02，density_grid_size=2500。"],
            ["深度学习 teacher", "MLPRegressor hidden_layer_sizes=(48,24)，默认 5 个 seed 集成。"],
            ["符号搜索", "time_power 从 -3.0 到 -0.25，默认步长 0.05；charge_power 在 -2 到 2 的半整数网格。"],
            ["结果版本", "聚类 q-ai-clustering-v8；符号回归 q-ai-clustering-symbolic-v8。"],
            ["近期稳定性修复", "只对非空 ChargeCluster 做整数转换，避免 NaN/inf 转 int 失败。"],
        ],
        [2300, 7060],
        header_fill=LIGHT_BLUE,
    )

    add_para(doc, "本报告基于当前本地代码和最近一次调试后的实现状态撰写。若后续继续调整模型参数、页面流程或结果字段，应同步更新报告中的参数说明和教学建议。")

    doc.core_properties.title = "密立根油滴 AI 实验平台分析报告"
    doc.core_properties.subject = "AI 聚类、深度学习 teacher、符号回归与物理实验教学"
    doc.core_properties.author = "Codex"
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build()
